import re
from dataclasses import dataclass
from io import BytesIO
from typing import Any

from fastapi import HTTPException, status
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.ai.schemas import PlaybookReviewOutput
from app.auth.models import User
from app.contract_files.models import (
    ContractEdit,
    ContractFile,
    ContractTextSnapshot,
    ContractVersion,
    StorageObject,
)
from app.contract_files.service import next_version_number
from app.contracts.models import Contract
from app.core.audit import write_audit_log, write_timeline_event
from app.core.enums import ContractVersionSource, PlaybookStatus, StorageBackend
from app.integrations.storage import storage_service
from app.playbooks.models import (
    Playbook,
    PlaybookDecision,
    PlaybookDeviation,
    PlaybookRule,
    PlaybookRun,
    PlaybookVersion,
)


GENERATED_RULE_TEMPLATES = [
    {
        "clause_type": "confidentiality",
        "rule_type": "required_language",
        "required_language": "confidential information",
        "preferred_position": "The contract should define confidential information, exclusions, permitted disclosures, and survival.",
        "risk_level": "medium",
        "rationale": "Confidentiality terms should be explicit enough to enforce and operationalize.",
        "sample_clause": "Each party shall protect the other party's confidential information using reasonable care.",
        "negotiation_guidance": "Accept mutual confidentiality where both parties exchange sensitive information.",
    },
    {
        "clause_type": "limitation_of_liability",
        "rule_type": "prohibited_language",
        "prohibited_language": "unlimited liability",
        "preferred_position": "Liability should be capped except for negotiated carve-outs.",
        "fallback_position": "Escalate uncapped liability to legal approval if business needs require it.",
        "risk_level": "high",
        "approval_required": True,
        "rationale": "Uncapped liability can create disproportionate exposure.",
        "negotiation_guidance": "Prefer a fees-based cap with narrow exclusions.",
    },
    {
        "clause_type": "termination",
        "rule_type": "required_language",
        "required_language": "notice",
        "preferred_position": "Termination rights should include notice periods, cure rights when appropriate, and effects of termination.",
        "risk_level": "medium",
        "rationale": "Termination mechanics reduce operational and dispute risk.",
        "negotiation_guidance": "Ask for a cure period for remediable breaches.",
    },
    {
        "clause_type": "governing_law",
        "rule_type": "required_language",
        "required_language": "governing law",
        "preferred_position": "The contract should state governing law and dispute forum.",
        "risk_level": "medium",
        "rationale": "Missing governing law increases dispute uncertainty.",
        "negotiation_guidance": "Prefer the organization's standard jurisdiction unless business context requires otherwise.",
    },
    {
        "clause_type": "assignment",
        "rule_type": "required_language",
        "required_language": "assignment",
        "preferred_position": "Assignments should require consent, with reasonable exceptions for affiliates and M&A transactions.",
        "risk_level": "low",
        "rationale": "Assignment controls prevent unexpected counterparty changes.",
        "negotiation_guidance": "Allow assignment to affiliates or successors in interest when acceptable.",
    },
]


@dataclass(frozen=True)
class EvaluatedDeviation:
    rule: PlaybookRule | None
    severity: str
    clause_type: str
    issue: str
    suggested_fix: str | None
    citation: dict[str, Any] | None
    original_text: str | None
    replacement_text: str | None


@dataclass(frozen=True)
class PlaybookRunArtifacts:
    run: PlaybookRun
    deviations: list[PlaybookDeviation]
    redline_version: ContractVersion | None
    contract_edit: ContractEdit | None


def get_playbook_for_user(db: Session, *, playbook_id: str, user: User) -> Playbook:
    playbook = db.get(Playbook, playbook_id)
    if playbook is None or playbook.org_id != user.org_id or playbook.deleted_at is not None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Playbook not found")
    return playbook


def get_playbook_version(
    db: Session,
    *,
    playbook: Playbook,
    version_id: str,
    org_id: str,
) -> PlaybookVersion:
    version = db.get(PlaybookVersion, version_id)
    if version is None or version.org_id != org_id or version.playbook_id != playbook.id:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Playbook version not found")
    return version


def create_initial_playbook(
    db: Session,
    *,
    user: User,
    name: str,
    description: str | None,
    generated_rules: list[dict[str, Any]] | None = None,
) -> Playbook:
    playbook = Playbook(
        org_id=user.org_id,
        name=name,
        description=description,
        status=PlaybookStatus.DRAFT,
        created_by_user_id=user.id,
        updated_by_user_id=user.id,
    )
    db.add(playbook)
    db.flush()
    version = PlaybookVersion(
        org_id=user.org_id,
        playbook_id=playbook.id,
        version_number=1,
        status=PlaybookStatus.DRAFT,
        summary="Initial draft",
        source_metadata={"generated": bool(generated_rules)},
        created_by_user_id=user.id,
        updated_by_user_id=user.id,
    )
    db.add(version)
    db.flush()
    playbook.current_version_id = version.id
    for rule_payload in generated_rules or []:
        db.add(
            PlaybookRule(
                org_id=user.org_id,
                playbook_version_id=version.id,
                created_by_user_id=user.id,
                updated_by_user_id=user.id,
                **rule_payload,
            )
        )
    write_audit_log(
        db,
        action="playbook.created",
        resource_type="playbook",
        resource_id=playbook.id,
        org_id=user.org_id,
        actor_user_id=user.id,
        after={"name": playbook.name, "generated_rule_count": len(generated_rules or [])},
    )
    return playbook


def generated_default_rules(*, contract_type: str | None, focus_areas: list[str]) -> list[dict[str, Any]]:
    requested = {area.strip().lower().replace(" ", "_") for area in focus_areas if area.strip()}
    rules: list[dict[str, Any]] = []
    for template in GENERATED_RULE_TEMPLATES:
        if requested and template["clause_type"] not in requested:
            continue
        rule = dict(template)
        rule.setdefault("fallback_position", None)
        rule.setdefault("prohibited_language", None)
        rule.setdefault("required_language", None)
        rule.setdefault("approval_required", False)
        rule.setdefault("escalation_role", "legal")
        rule.setdefault("sample_clause", None)
        rule.setdefault("negotiation_guidance", None)
        if contract_type:
            rule["rationale"] = f"{rule['rationale']} Generated for {contract_type} playbook use."
        rules.append(rule)
    return rules or [dict(rule) for rule in GENERATED_RULE_TEMPLATES]


def next_playbook_version_number(db: Session, *, playbook_id: str) -> int:
    max_version = db.scalar(
        select(func.max(PlaybookVersion.version_number)).where(PlaybookVersion.playbook_id == playbook_id)
    )
    return int(max_version or 0) + 1


def clone_playbook_version(
    db: Session,
    *,
    playbook: Playbook,
    user: User,
    source_version: PlaybookVersion | None,
    summary: str | None,
) -> PlaybookVersion:
    version = PlaybookVersion(
        org_id=user.org_id,
        playbook_id=playbook.id,
        version_number=next_playbook_version_number(db, playbook_id=playbook.id),
        status=PlaybookStatus.DRAFT,
        summary=summary or f"Draft version cloned from V{source_version.version_number}" if source_version else summary,
        source_metadata={"source_version_id": source_version.id if source_version else None},
        created_by_user_id=user.id,
        updated_by_user_id=user.id,
    )
    db.add(version)
    db.flush()
    if source_version is not None:
        source_rules = db.scalars(
            select(PlaybookRule).where(
                PlaybookRule.org_id == user.org_id,
                PlaybookRule.playbook_version_id == source_version.id,
            )
        ).all()
        for rule in source_rules:
            db.add(
                PlaybookRule(
                    org_id=user.org_id,
                    playbook_version_id=version.id,
                    clause_type=rule.clause_type,
                    rule_type=rule.rule_type,
                    preferred_position=rule.preferred_position,
                    fallback_position=rule.fallback_position,
                    prohibited_language=rule.prohibited_language,
                    required_language=rule.required_language,
                    risk_level=rule.risk_level,
                    rationale=rule.rationale,
                    approval_required=rule.approval_required,
                    escalation_role=rule.escalation_role,
                    sample_clause=rule.sample_clause,
                    negotiation_guidance=rule.negotiation_guidance,
                    created_by_user_id=user.id,
                    updated_by_user_id=user.id,
                )
            )
    playbook.current_version_id = version.id
    playbook.updated_by_user_id = user.id
    write_audit_log(
        db,
        action="playbook.version_created",
        resource_type="playbook_version",
        resource_id=version.id,
        org_id=user.org_id,
        actor_user_id=user.id,
        after={"playbook_id": playbook.id, "source_version_id": source_version.id if source_version else None},
    )
    return version


def select_run_version(
    db: Session,
    *,
    playbook: Playbook,
    org_id: str,
    version_id: str | None,
    test_mode: bool,
) -> PlaybookVersion:
    if version_id:
        version = get_playbook_version(db, playbook=playbook, version_id=version_id, org_id=org_id)
    else:
        version = db.scalar(
            select(PlaybookVersion)
            .where(
                PlaybookVersion.org_id == org_id,
                PlaybookVersion.playbook_id == playbook.id,
                PlaybookVersion.status == PlaybookStatus.PUBLISHED,
            )
            .order_by(PlaybookVersion.version_number.desc())
        )
        if version is None and test_mode and playbook.current_version_id:
            version = get_playbook_version(
                db, playbook=playbook, version_id=playbook.current_version_id, org_id=org_id
            )
    if version is None:
        raise HTTPException(status.HTTP_409_CONFLICT, "No published playbook version is available")
    if version.status != PlaybookStatus.PUBLISHED and not test_mode:
        raise HTTPException(status.HTTP_409_CONFLICT, "Only published playbook versions can be run officially")
    return version


def evaluate_rules_against_text(*, rules: list[PlaybookRule], text: str) -> list[EvaluatedDeviation]:
    deviations: list[EvaluatedDeviation] = []
    for rule in rules:
        prohibited = _clean_phrase(rule.prohibited_language)
        if prohibited:
            match = _find_phrase(text, prohibited)
            if match:
                start, end = match
                quote = text[start:end]
                deviations.append(
                    EvaluatedDeviation(
                        rule=rule,
                        severity=_normalize_severity(rule.risk_level, approval_required=rule.approval_required),
                        clause_type=rule.clause_type,
                        issue=f"Prohibited language found for {rule.clause_type}: {prohibited}",
                        suggested_fix=_suggested_fix(rule),
                        citation={
                            "type": "text_match",
                            "quote": quote,
                            "start_char": start,
                            "end_char": end,
                        },
                        original_text=quote,
                        replacement_text=_suggested_fix(rule),
                    )
                )
        required = _clean_phrase(rule.required_language)
        if required and _find_phrase(text, required) is None:
            deviations.append(
                EvaluatedDeviation(
                    rule=rule,
                    severity=_normalize_severity(rule.risk_level, approval_required=rule.approval_required),
                    clause_type=rule.clause_type,
                    issue=f"Required language missing for {rule.clause_type}: {required}",
                    suggested_fix=_suggested_fix(rule),
                    citation={
                        "type": "playbook_rule",
                        "rule_id": rule.id,
                        "required_language": required,
                    },
                    original_text=None,
                    replacement_text=_suggested_fix(rule),
                )
            )
    return deviations


def ai_deviations_to_evaluated(
    *,
    ai_output: PlaybookReviewOutput | None,
    rules: list[PlaybookRule],
) -> list[EvaluatedDeviation]:
    if ai_output is None or not ai_output.deviations:
        return []
    by_index = {index + 1: rule for index, rule in enumerate(rules)}
    evaluated: list[EvaluatedDeviation] = []
    for deviation in ai_output.deviations:
        rule = by_index.get(deviation.rule_index) if deviation.rule_index is not None else None
        citation = None
        if deviation.citations:
            first = deviation.citations[0]
            citation = {
                "type": "ai_citation",
                "quote": first.quote,
                "start_char": first.start_char,
                "end_char": first.end_char,
                "page_number": first.page_number,
            }
        evaluated.append(
            EvaluatedDeviation(
                rule=rule,
                severity=deviation.severity,
                clause_type=deviation.clause_type or (rule.clause_type if rule else "unknown"),
                issue=deviation.issue,
                suggested_fix=deviation.suggested_fix,
                citation=citation,
                original_text=deviation.original_text,
                replacement_text=deviation.suggested_fix,
            )
        )
    return evaluated


def execute_playbook_run(
    db: Session,
    *,
    user: User,
    playbook: Playbook,
    version: PlaybookVersion,
    contract: Contract,
    create_redline: bool,
    ai_output: PlaybookReviewOutput | None = None,
    ai_error: str | None = None,
) -> PlaybookRunArtifacts:
    contract_version, contract_file, snapshot = _current_contract_artifacts(db, contract=contract, org_id=user.org_id)
    rules = list(
        db.scalars(
            select(PlaybookRule).where(
                PlaybookRule.org_id == user.org_id,
                PlaybookRule.playbook_version_id == version.id,
            )
        ).all()
    )
    if not rules:
        raise HTTPException(status.HTTP_409_CONFLICT, "Playbook version has no rules")
    run = PlaybookRun(
        org_id=user.org_id,
        playbook_id=playbook.id,
        playbook_version_id=version.id,
        contract_id=contract.id,
        contract_version_id=contract_version.id,
        status="running",
        raw_ai_output=ai_output.model_dump(mode="json") if ai_output else {"mode": "deterministic_playbook_engine"},
        validation_status="not_validated",
        model_name="claude" if ai_output else "deterministic_playbook_engine",
        error_message=ai_error,
        created_by_user_id=user.id,
        updated_by_user_id=user.id,
    )
    db.add(run)
    db.flush()
    evaluated = ai_deviations_to_evaluated(ai_output=ai_output, rules=rules)
    deviation_source = "claude"
    if not evaluated:
        evaluated = evaluate_rules_against_text(rules=rules, text=snapshot.text)
        deviation_source = "deterministic"
    deviations = [
        _create_deviation(
            db,
            user=user,
            run=run,
            contract=contract,
            evaluated=deviation,
        )
        for deviation in evaluated
    ]
    redline_version = None
    contract_edit = None
    if create_redline and deviations:
        redline_version, contract_edit = _create_playbook_redline_version(
            db,
            user=user,
            contract=contract,
            contract_file=contract_file,
            base_version=contract_version,
            source_snapshot=snapshot,
            run=run,
            deviations=deviations,
        )
    run.status = "succeeded"
    run.validation_status = "valid"
    run.validated_output = {
        "deviation_count": len(deviations),
        "deviation_source": deviation_source,
        "redline_contract_version_id": redline_version.id if redline_version else None,
        "contract_edit_id": contract_edit.id if contract_edit else None,
    }
    write_audit_log(
        db,
        action="playbook.run_completed",
        resource_type="playbook_run",
        resource_id=run.id,
        org_id=user.org_id,
        actor_user_id=user.id,
        after={
            "playbook_id": playbook.id,
            "playbook_version_id": version.id,
            "contract_id": contract.id,
            "deviation_count": len(deviations),
            "redline_contract_version_id": redline_version.id if redline_version else None,
        },
    )
    write_timeline_event(
        db,
        org_id=user.org_id,
        resource_type="contract",
        resource_id=contract.id,
        event_type="playbook.run_completed",
        title="Playbook review completed",
        actor_user_id=user.id,
        details={
            "playbook_id": playbook.id,
            "playbook_name": playbook.name,
            "playbook_run_id": run.id,
            "deviation_count": len(deviations),
            "redline_contract_version_id": redline_version.id if redline_version else None,
        },
    )
    db.flush()
    return PlaybookRunArtifacts(run=run, deviations=deviations, redline_version=redline_version, contract_edit=contract_edit)


def record_deviation_decision(
    db: Session,
    *,
    deviation: PlaybookDeviation,
    decision: str,
    rationale: str | None,
    user: User,
) -> PlaybookDecision:
    row = PlaybookDecision(
        org_id=user.org_id,
        playbook_deviation_id=deviation.id,
        decision=decision,
        rationale=rationale,
        decided_by_user_id=user.id,
        created_by_user_id=user.id,
        updated_by_user_id=user.id,
    )
    deviation.status = decision
    deviation.updated_by_user_id = user.id
    db.add(row)
    db.flush()
    write_audit_log(
        db,
        action="playbook.deviation_decided",
        resource_type="playbook_deviation",
        resource_id=deviation.id,
        org_id=user.org_id,
        actor_user_id=user.id,
        after={"decision": decision, "rationale": rationale},
    )
    write_timeline_event(
        db,
        org_id=user.org_id,
        resource_type="contract",
        resource_id=deviation.contract_id,
        event_type="playbook.deviation_decided",
        title="Playbook deviation decision recorded",
        actor_user_id=user.id,
        details={"playbook_deviation_id": deviation.id, "decision": decision},
    )
    return row


def _create_deviation(
    db: Session,
    *,
    user: User,
    run: PlaybookRun,
    contract: Contract,
    evaluated: EvaluatedDeviation,
) -> PlaybookDeviation:
    row = PlaybookDeviation(
        org_id=user.org_id,
        playbook_run_id=run.id,
        playbook_rule_id=evaluated.rule.id if evaluated.rule else None,
        contract_id=contract.id,
        severity=evaluated.severity,
        clause_type=evaluated.clause_type,
        issue=evaluated.issue,
        suggested_fix=evaluated.suggested_fix,
        citation=evaluated.citation,
        status="open",
        created_by_user_id=user.id,
        updated_by_user_id=user.id,
    )
    db.add(row)
    db.flush()
    return row


def _current_contract_artifacts(
    db: Session,
    *,
    contract: Contract,
    org_id: str,
) -> tuple[ContractVersion, ContractFile, ContractTextSnapshot]:
    version = db.get(ContractVersion, contract.current_authoritative_version_id) if contract.current_authoritative_version_id else None
    if version is None or version.org_id != org_id or version.contract_id != contract.id:
        raise HTTPException(status.HTTP_409_CONFLICT, "Contract has no authoritative version")
    contract_file = db.get(ContractFile, version.contract_file_id)
    if contract_file is None or contract_file.org_id != org_id:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Contract file not found")
    snapshot = db.get(ContractTextSnapshot, version.text_snapshot_id) if version.text_snapshot_id else None
    if snapshot is None or snapshot.org_id != org_id or snapshot.contract_id != contract.id:
        raise HTTPException(status.HTTP_409_CONFLICT, "Contract has no text snapshot to review")
    return version, contract_file, snapshot


def _create_playbook_redline_version(
    db: Session,
    *,
    user: User,
    contract: Contract,
    contract_file: ContractFile,
    base_version: ContractVersion,
    source_snapshot: ContractTextSnapshot,
    run: PlaybookRun,
    deviations: list[PlaybookDeviation],
) -> tuple[ContractVersion, ContractEdit]:
    instructions = _redline_instructions(deviations)
    proposed_text = _proposed_text(source_snapshot.text, instructions)
    content = _build_playbook_redline_docx(
        contract_title=contract.title,
        base_version_number=base_version.version_number,
        source_text=source_snapshot.text,
        instructions=instructions,
    )
    storage_object = _store_docx(
        db,
        org_id=user.org_id,
        user_id=user.id,
        filename=f"{_safe_filename(contract.title)}-playbook-redline-v{base_version.version_number}.docx",
        content=content,
    )
    edit = ContractEdit(
        org_id=user.org_id,
        contract_id=contract.id,
        contract_version_id=base_version.id,
        edit_type="playbook_redline",
        status="proposed",
        original_text=source_snapshot.text[:4000],
        replacement_text=proposed_text,
        rationale=f"Playbook run {run.id} produced {len(deviations)} deviation(s).",
        citation=[
            {
                "type": "playbook_run",
                "playbook_run_id": run.id,
                "playbook_deviation_ids": [row.id for row in deviations],
            }
        ],
        created_by_user_id=user.id,
        updated_by_user_id=user.id,
    )
    db.add(edit)
    db.flush()
    redline_version = ContractVersion(
        org_id=user.org_id,
        contract_id=contract.id,
        contract_file_id=contract_file.id,
        version_number=next_version_number(db, contract_file.id),
        storage_object_id=storage_object.id,
        source=ContractVersionSource.PLAYBOOK_REDLINE,
        change_summary=f"Playbook redline proposal from run {run.id}",
        is_authoritative=False,
        created_by_user_id=user.id,
        updated_by_user_id=user.id,
    )
    db.add(redline_version)
    db.flush()
    edit.citation = [
        *(edit.citation or []),
        {"type": "assistant_edit_version", "contract_version_id": redline_version.id},
    ]
    snapshot = ContractTextSnapshot(
        org_id=user.org_id,
        contract_id=contract.id,
        contract_version_id=redline_version.id,
        extraction_method="playbook_redline_text",
        extraction_quality_score=1.0,
        text=proposed_text,
        page_map=source_snapshot.page_map,
        ocr_provider=source_snapshot.ocr_provider,
        validation_status="complete",
        created_by_user_id=user.id,
        updated_by_user_id=user.id,
    )
    db.add(snapshot)
    db.flush()
    redline_version.text_snapshot_id = snapshot.id
    write_audit_log(
        db,
        action="playbook.redline_created",
        resource_type="contract_version",
        resource_id=redline_version.id,
        org_id=user.org_id,
        actor_user_id=user.id,
        after={"playbook_run_id": run.id, "contract_edit_id": edit.id, "base_version_id": base_version.id},
    )
    write_timeline_event(
        db,
        org_id=user.org_id,
        resource_type="contract",
        resource_id=contract.id,
        event_type="playbook.redline_created",
        title="Playbook tracked-change redline created",
        actor_user_id=user.id,
        details={"playbook_run_id": run.id, "contract_version_id": redline_version.id, "contract_edit_id": edit.id},
    )
    return redline_version, edit


def _store_docx(db: Session, *, org_id: str, user_id: str, filename: str, content: bytes) -> StorageObject:
    stored = storage_service.save_bytes(
        org_id=org_id,
        filename=filename,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        content=content,
    )
    storage_object = StorageObject(
        org_id=org_id,
        storage_key=stored.storage_key,
        filename=stored.filename,
        mime_type=stored.mime_type,
        size_bytes=stored.size_bytes,
        sha256_hash=stored.sha256_hash,
        storage_backend=StorageBackend.LOCAL_VOLUME,
        created_by_user_id=user_id,
        updated_by_user_id=user_id,
    )
    db.add(storage_object)
    db.flush()
    return storage_object


def _build_playbook_redline_docx(
    *,
    contract_title: str,
    base_version_number: int,
    source_text: str,
    instructions: str,
) -> bytes:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = Document()
    _enable_word_track_revisions(document, OxmlElement=OxmlElement, qn=qn)
    document.add_heading(f"{contract_title} - Playbook Redline Proposal", level=1)
    document.add_paragraph(f"Base version: V{base_version_number}")
    document.add_heading("Playbook Findings", level=2)
    document.add_paragraph(instructions)
    document.add_heading("Native Word Tracked Changes", level=2)
    revision_paragraph = document.add_paragraph()
    _append_deleted_text(
        revision_paragraph,
        source_text or "No source text snapshot was available.",
        author="Legal AI Playbook",
        revision_id="1",
        OxmlElement=OxmlElement,
        qn=qn,
    )
    _append_inserted_text(
        revision_paragraph,
        _proposed_text(source_text, instructions),
        author="Legal AI Playbook",
        revision_id="2",
        OxmlElement=OxmlElement,
        qn=qn,
    )
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _enable_word_track_revisions(document: Any, *, OxmlElement: Any, qn: Any) -> None:
    settings = document.settings.element
    if settings.find(qn("w:trackRevisions")) is None:
        settings.append(OxmlElement("w:trackRevisions"))


def _append_deleted_text(
    paragraph: Any,
    text: str,
    *,
    author: str,
    revision_id: str,
    OxmlElement: Any,
    qn: Any,
) -> None:
    deletion = OxmlElement("w:del")
    _set_revision_attrs(deletion, author=author, revision_id=revision_id, qn=qn)
    _append_revision_runs(deletion, text, text_tag="w:delText", OxmlElement=OxmlElement)
    paragraph._p.append(deletion)


def _append_inserted_text(
    paragraph: Any,
    text: str,
    *,
    author: str,
    revision_id: str,
    OxmlElement: Any,
    qn: Any,
) -> None:
    insertion = OxmlElement("w:ins")
    _set_revision_attrs(insertion, author=author, revision_id=revision_id, qn=qn)
    _append_revision_runs(insertion, text, text_tag="w:t", OxmlElement=OxmlElement)
    paragraph._p.append(insertion)


def _set_revision_attrs(element: Any, *, author: str, revision_id: str, qn: Any) -> None:
    from app.core.database import utcnow

    element.set(qn("w:id"), revision_id)
    element.set(qn("w:author"), author)
    element.set(qn("w:date"), utcnow().replace(microsecond=0).isoformat())


def _append_revision_runs(parent: Any, text: str, *, text_tag: str, OxmlElement: Any) -> None:
    xml_space = "{http://www.w3.org/XML/1998/namespace}space"
    for line_index, line in enumerate(text.splitlines() or [""]):
        if line_index:
            break_run = OxmlElement("w:r")
            break_run.append(OxmlElement("w:br"))
            parent.append(break_run)
        run = OxmlElement("w:r")
        text_element = OxmlElement(text_tag)
        text_element.set(xml_space, "preserve")
        text_element.text = line
        run.append(text_element)
        parent.append(run)


def _proposed_text(source_text: str, instructions: str) -> str:
    return "\n\n".join([source_text, "[Playbook redline suggestions]", instructions]).strip()


def _redline_instructions(deviations: list[PlaybookDeviation]) -> str:
    lines = []
    for index, deviation in enumerate(deviations, start=1):
        fix = deviation.suggested_fix or "Review and revise to match the playbook position."
        lines.append(f"{index}. {deviation.clause_type} ({deviation.severity}): {deviation.issue}\nSuggested fix: {fix}")
    return "\n\n".join(lines)


def _clean_phrase(value: str | None) -> str | None:
    if value is None:
        return None
    cleaned = " ".join(value.split())
    return cleaned or None


def _find_phrase(text: str, phrase: str) -> tuple[int, int] | None:
    normalized = re.escape(" ".join(phrase.split()))
    pattern = re.compile(normalized.replace(r"\ ", r"\s+"), flags=re.IGNORECASE)
    match = pattern.search(text)
    if match is None:
        return None
    return match.start(), match.end()


def _suggested_fix(rule: PlaybookRule) -> str | None:
    return rule.preferred_position or rule.fallback_position or rule.required_language or rule.sample_clause


def _normalize_severity(value: str | None, *, approval_required: bool) -> str:
    normalized = (value or "medium").lower()
    if normalized not in {"low", "medium", "high", "critical"}:
        normalized = "medium"
    if approval_required and normalized in {"low", "medium"}:
        return "high"
    return normalized


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", value).strip("-._")
    return cleaned[:120] or "contract"
