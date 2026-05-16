import hashlib
import json
import logging
import re
import secrets
from datetime import UTC, datetime, timedelta
from io import BytesIO
from typing import Any

from fastapi import HTTPException, status
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.ai.confirmations import create_confirmation
from app.ai.models import AIConfirmation
from app.ai.tool_registry import (
    ApprovalSubmitInput,
    ArchiveContractInput,
    BrainAskInput,
    ContractHandleInput,
    EditContractInput,
    ExternalShareInput,
    ExtractObligationsInput,
    FindInContractInput,
    GenerateContractInput,
    PlaybookToolInput,
    ProjectContractsInput,
    ReadTableCellsInput,
    SignatureSendInput,
    TabularReviewInput,
    WorkflowRunInput,
    tool_registry,
)
from app.ai.schemas import BrainQueryParseOutput
from app.assistant.models import AssistantContractHandle, AssistantToolCall
from app.auth.models import User
from app.approvals.service import submit_contract_for_approval
from app.contract_brain.retrieval import assemble_context
from app.contract_files.models import (
    ContractShare,
    ContractEdit,
    ContractFile,
    ContractTextSnapshot,
    ContractVersion,
    StorageObject,
)
from app.contract_files.service import _queue_initial_contract_jobs, next_version_number
from app.contracts.models import Contract
from app.contracts.lifecycle import transition_contract_stage
from app.contracts.service import get_contract_for_user
from app.core.audit import write_audit_log, write_timeline_event
from app.jobs.models import JobRun
from app.jobs.service import create_job, dispatch_job
from app.core.database import utcnow
from app.core.enums import (
    AssistantToolCallStatus,
    ContractLifecycleStage,
    ContractVersionSource,
    ShareAccessMode,
    SignatureStatus,
    StorageBackend,
    TabularCellStatus,
)
from app.core.rbac import has_permission
from app.integrations.docusign import docusign_client
from app.integrations.resend import resend_client
from app.integrations.storage import storage_service
from app.playbooks.service import execute_playbook_run, get_playbook_for_user, select_run_version
from app.projects.access import get_project_for_user
from app.projects.models import ProjectContract
from app.signatures.models import SignatureRecipient, SignatureRequest
from app.signatures.service import validate_signature_recipients
from app.tabular_review.models import TabularReview, TabularReviewCell, TabularReviewColumn
from app.tabular_review.service import dispatch_cells
from app.workflows.models import Workflow, WorkflowRun


class ToolRuntime:
    async def execute(
        self,
        db: Session,
        *,
        tool_name: str,
        tool_input: dict[str, Any],
        user: User,
        session_id: str,
        assistant_run_id: str | None = None,
        provider_tool_use_id: str | None = None,
        provider_state: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        spec = tool_registry.get(tool_name)
        if not has_permission(user.permission_values, spec.required_permission):
            raise HTTPException(status.HTTP_403_FORBIDDEN, f"Missing permission: {spec.required_permission}")
        validated_input = spec.input_model.model_validate(tool_input)
        call = AssistantToolCall(
            org_id=user.org_id,
            session_id=session_id,
            assistant_run_id=assistant_run_id,
            provider_tool_use_id=provider_tool_use_id,
            tool_name=tool_name,
            category=spec.category,
            arguments=validated_input.model_dump(mode="json"),
            status=AssistantToolCallStatus.RUNNING,
            confirmation_required=spec.requires_confirmation,
            idempotency_key=_idempotency_key(tool_name, session_id, validated_input.model_dump(mode="json")),
            output_schema_name=spec.output_model.__name__,
            started_at=utcnow(),
            created_by_user_id=user.id,
            updated_by_user_id=user.id,
        )
        db.add(call)
        db.flush()
        if spec.requires_confirmation:
            confirmation = create_confirmation(
                db,
                org_id=user.org_id,
                user_id=user.id,
                session_id=session_id,
                assistant_run_id=assistant_run_id or "",
                tool_call=call,
                tool_input=validated_input.model_dump(mode="json"),
                policy={"confirmation_policy": spec.confirmation_policy, "category": spec.category},
                provider_state=provider_state,
            )
            db.flush()
            return {
                "confirmation_required": True,
                "tool_call_id": call.id,
                "confirmation_id": confirmation.id,
            }

        try:
            result = await self._execute_validated(
                db,
                tool_name=tool_name,
                payload=validated_input,
                user=user,
                session_id=session_id,
            )
            call.result = result
            call.status = AssistantToolCallStatus.SUCCEEDED
            call.finished_at = utcnow()
            db.flush()
            return result
        except Exception as exc:
            call.status = AssistantToolCallStatus.FAILED
            call.error_message = str(exc)
            call.finished_at = utcnow()
            db.flush()
            raise

    async def execute_confirmed(
        self,
        db: Session,
        *,
        confirmation: AIConfirmation,
        user: User,
    ) -> dict[str, Any]:
        call = db.get(AssistantToolCall, confirmation.tool_call_id)
        if call is None or call.org_id != user.org_id:
            raise HTTPException(status.HTTP_404_NOT_FOUND, "Tool call not found")
        spec = tool_registry.get(call.tool_name)
        if not has_permission(user.permission_values, spec.required_permission):
            raise HTTPException(status.HTTP_403_FORBIDDEN, f"Missing permission: {spec.required_permission}")
        if call.status not in {AssistantToolCallStatus.CONFIRMED, AssistantToolCallStatus.RUNNING}:
            raise HTTPException(status.HTTP_409_CONFLICT, "Tool call is not confirmed")
        validated_input = spec.input_model.model_validate(call.arguments)
        call.status = AssistantToolCallStatus.RUNNING
        call.started_at = call.started_at or utcnow()
        try:
            result = await self._execute_validated(
                db,
                tool_name=call.tool_name,
                payload=validated_input,
                user=user,
                session_id=call.session_id,
            )
            call.result = result
            call.status = AssistantToolCallStatus.SUCCEEDED
            call.finished_at = utcnow()
            call.updated_by_user_id = user.id
            db.flush()
            return result
        except Exception as exc:
            call.status = AssistantToolCallStatus.FAILED
            call.error_message = str(exc)
            call.finished_at = utcnow()
            call.updated_by_user_id = user.id
            db.flush()
            raise

    async def _execute_validated(
        self,
        db: Session,
        *,
        tool_name: str,
        payload: Any,
        user: User,
        session_id: str,
    ) -> dict[str, Any]:
        if tool_name == "read_contract":
            return self._read_contract(db, payload=payload, user=user, session_id=session_id)
        if tool_name == "find_in_contract":
            return self._find_in_contract(db, payload=payload, user=user, session_id=session_id)
        if tool_name == "list_project_contracts":
            return self._list_project_contracts(db, payload=payload, user=user)
        if tool_name == "get_contract_status":
            return self._get_contract_status(db, payload=payload, user=user, session_id=session_id)
        if tool_name == "list_workflows":
            return self._list_workflows(db, user=user)
        if tool_name == "run_workflow":
            return self._run_workflow(db, payload=payload, user=user)
        if tool_name == "generate_contract_docx":
            return self._generate_contract_docx(db, payload=payload, user=user)
        if tool_name == "edit_contract":
            return self._edit_contract(db, payload=payload, user=user, session_id=session_id)
        if tool_name == "replicate_contract_version":
            return self._replicate_contract_version(db, payload=payload, user=user, session_id=session_id)
        if tool_name == "run_playbook_review":
            return self._run_playbook_review(
                db,
                payload=payload,
                user=user,
                session_id=session_id,
                create_redline=False,
            )
        if tool_name == "redline_against_playbook":
            return self._run_playbook_review(
                db,
                payload=payload,
                user=user,
                session_id=session_id,
                create_redline=True,
            )
        if tool_name == "ask_contract_brain":
            return self._ask_contract_brain(db, payload=payload, user=user, session_id=session_id)
        if tool_name == "submit_for_approval":
            return await self._submit_for_approval(db, payload=payload, user=user, session_id=session_id)
        if tool_name == "send_for_signature":
            return await self._send_for_signature(db, payload=payload, user=user, session_id=session_id)
        if tool_name == "extract_obligations":
            return self._extract_obligations(db, payload=payload, user=user, session_id=session_id)
        if tool_name == "create_tabular_review":
            return self._create_tabular_review(db, payload=payload, user=user, session_id=session_id)
        if tool_name == "read_table_cells":
            return self._read_table_cells(db, payload=payload, user=user)
        if tool_name == "external_share":
            return self._external_share(db, payload=payload, user=user, session_id=session_id)
        if tool_name == "archive_contract":
            return self._archive_contract(db, payload=payload, user=user, session_id=session_id)
        return {"status": "feature_not_enabled", "tool": tool_name}

    def _read_contract(
        self,
        db: Session,
        *,
        payload: ContractHandleInput,
        user: User,
        session_id: str,
    ) -> dict[str, Any]:
        contract = self._resolve_contract(db, payload=payload, user=user, session_id=session_id)
        version = db.get(ContractVersion, contract.current_authoritative_version_id) if contract.current_authoritative_version_id else None
        snapshot = db.get(ContractTextSnapshot, version.text_snapshot_id) if version and version.text_snapshot_id else None
        return {
            "contract_id": contract.id,
            "title": contract.title,
            "lifecycle_stage": contract.lifecycle_stage,
            "text_snapshot_id": snapshot.id if snapshot else None,
            "text_excerpt": (snapshot.text[:8000] if snapshot else ""),
            "text_truncated": bool(snapshot and len(snapshot.text) > 8000),
        }

    def _find_in_contract(
        self,
        db: Session,
        *,
        payload: FindInContractInput,
        user: User,
        session_id: str,
    ) -> dict[str, Any]:
        contract = self._resolve_contract(db, payload=payload, user=user, session_id=session_id)
        version = db.get(ContractVersion, contract.current_authoritative_version_id) if contract.current_authoritative_version_id else None
        snapshot = db.get(ContractTextSnapshot, version.text_snapshot_id) if version and version.text_snapshot_id else None
        text = snapshot.text if snapshot else ""
        matches = []
        lower_text = text.lower()
        query = payload.query.lower()
        start = lower_text.find(query)
        while start >= 0 and len(matches) < 10:
            end = start + len(query)
            excerpt_start = max(0, start - 160)
            excerpt_end = min(len(text), end + 160)
            matches.append({"start_char": start, "end_char": end, "excerpt": text[excerpt_start:excerpt_end]})
            start = lower_text.find(query, end)
        return {"contract_id": contract.id, "query": payload.query, "matches": matches}

    def _list_project_contracts(self, db: Session, *, payload: ProjectContractsInput, user: User) -> dict[str, Any]:
        get_project_for_user(db, project_id=payload.project_id, user=user)
        contract_ids = db.scalars(
            select(ProjectContract.contract_id).where(
                ProjectContract.org_id == user.org_id,
                ProjectContract.project_id == payload.project_id,
            )
        ).all()
        contracts = db.scalars(
            select(Contract).where(Contract.org_id == user.org_id, Contract.id.in_(contract_ids))
        ).all() if contract_ids else []
        return {
            "project_id": payload.project_id,
            "contracts": [
                {
                    "contract_id": contract.id,
                    "title": contract.title,
                    "lifecycle_stage": contract.lifecycle_stage,
                    "risk_level": contract.risk_level,
                }
                for contract in contracts
            ],
        }

    def _get_contract_status(
        self,
        db: Session,
        *,
        payload: ContractHandleInput,
        user: User,
        session_id: str,
    ) -> dict[str, Any]:
        contract = self._resolve_contract(db, payload=payload, user=user, session_id=session_id)
        return {
            "contract_id": contract.id,
            "title": contract.title,
            "lifecycle_stage": contract.lifecycle_stage,
            "risk_level": contract.risk_level,
            "counterparty_name": contract.counterparty_name,
            "current_authoritative_version_id": contract.current_authoritative_version_id,
        }

    def _list_workflows(self, db: Session, *, user: User) -> dict[str, Any]:
        workflows = db.scalars(
            select(Workflow).where(Workflow.org_id == user.org_id, Workflow.deleted_at.is_(None))
        ).all()
        return {"workflows": [{"workflow_id": row.id, "name": row.name, "workflow_type": row.workflow_type} for row in workflows]}

    def _run_workflow(self, db: Session, *, payload: WorkflowRunInput, user: User) -> dict[str, Any]:
        workflow = db.get(Workflow, payload.workflow_id)
        if workflow is None or workflow.org_id != user.org_id or workflow.deleted_at is not None:
            raise HTTPException(status.HTTP_404_NOT_FOUND, "Workflow not found")
        run = WorkflowRun(
            org_id=user.org_id,
            workflow_id=workflow.id,
            status="queued",
            input_contract_ids=payload.contract_ids,
            input_prompt=payload.prompt,
            created_by_user_id=user.id,
            updated_by_user_id=user.id,
        )
        db.add(run)
        db.flush()
        return {"workflow_run_id": run.id, "status": run.status}

    def _generate_contract_docx(
        self,
        db: Session,
        *,
        payload: GenerateContractInput,
        user: User,
    ) -> dict[str, Any]:
        if payload.project_id:
            get_project_for_user(db, project_id=payload.project_id, user=user, access="update")
        body_text, content = _build_generated_contract_docx(
            title=payload.title,
            instructions=payload.instructions,
        )
        storage_object = _store_docx(
            db,
            org_id=user.org_id,
            user_id=user.id,
            filename=f"{_safe_filename(payload.title)}.docx",
            content=content,
        )
        contract = Contract(
            org_id=user.org_id,
            title=payload.title,
            lifecycle_stage=ContractLifecycleStage.DRAFTING,
            owner_user_id=user.id,
            created_by_user_id=user.id,
            updated_by_user_id=user.id,
        )
        db.add(contract)
        db.flush()
        contract_file = ContractFile(
            org_id=user.org_id,
            contract_id=contract.id,
            file_label=storage_object.filename,
            created_by_user_id=user.id,
            updated_by_user_id=user.id,
        )
        db.add(contract_file)
        db.flush()
        version = ContractVersion(
            org_id=user.org_id,
            contract_id=contract.id,
            contract_file_id=contract_file.id,
            version_number=1,
            storage_object_id=storage_object.id,
            source=ContractVersionSource.ASSISTANT_GENERATED,
            change_summary="Assistant-generated DOCX",
            is_authoritative=True,
            created_by_user_id=user.id,
            updated_by_user_id=user.id,
        )
        db.add(version)
        db.flush()
        snapshot = ContractTextSnapshot(
            org_id=user.org_id,
            contract_id=contract.id,
            contract_version_id=version.id,
            extraction_method="assistant_generated_text",
            extraction_quality_score=1.0,
            text=body_text,
            page_map=None,
            validation_status="complete",
            created_by_user_id=user.id,
            updated_by_user_id=user.id,
        )
        db.add(snapshot)
        db.flush()
        version.text_snapshot_id = snapshot.id
        contract_file.current_version_id = version.id
        contract.current_contract_file_id = contract_file.id
        contract.current_authoritative_version_id = version.id
        if payload.project_id:
            db.add(
                ProjectContract(
                    org_id=user.org_id,
                    project_id=payload.project_id,
                    contract_id=contract.id,
                    created_by_user_id=user.id,
                    updated_by_user_id=user.id,
                )
            )
        write_audit_log(
            db,
            action="assistant.contract_generated",
            resource_type="contract",
            resource_id=contract.id,
            org_id=user.org_id,
            actor_user_id=user.id,
            after={
                "contract_file_id": contract_file.id,
                "contract_version_id": version.id,
                "storage_object_id": storage_object.id,
            },
        )
        write_timeline_event(
            db,
            org_id=user.org_id,
            resource_type="contract",
            resource_id=contract.id,
            event_type="assistant.contract_generated",
            title="Assistant generated contract",
            actor_user_id=user.id,
            details={"contract_version_id": version.id, "project_id": payload.project_id},
        )
        queued_jobs = _queue_initial_contract_jobs(
            db, user=user, contract=contract, version=version, snapshot=snapshot
        )
        db.flush()
        queued_job_ids = [job.id for job in queued_jobs]
        # Commit so the JobRun rows are visible before the Celery worker
        # picks them up (same ordering guarantee as the upload pipeline).
        db.commit()
        dispatched_job_types = []
        dispatch_errors = []
        for job_id in queued_job_ids:
            job = db.get(JobRun, job_id)
            if job is not None:
                try:
                    dispatch_job(db, job=job)
                    dispatched_job_types.append(job.job_type)
                except Exception as exc:
                    dispatch_errors.append(
                        {"job_id": job.id, "job_type": job.job_type, "error": str(exc)}
                    )
        if dispatched_job_types or dispatch_errors:
            write_timeline_event(
                db,
                org_id=user.org_id,
                resource_type="contract",
                resource_id=contract.id,
                event_type="contract.ai_jobs_dispatched",
                title="Contract AI jobs dispatched",
                actor_user_id=user.id,
                details={
                    "dispatched_job_types": dispatched_job_types,
                    "dispatch_errors": dispatch_errors,
                },
            )
        db.commit()
        return {
            "status": "created",
            "artifact_type": "generated_contract",
            "contract_id": contract.id,
            "contract_file_id": contract_file.id,
            "contract_version_id": version.id,
            "storage_object_id": storage_object.id,
        }

    def _edit_contract(
        self,
        db: Session,
        *,
        payload: EditContractInput,
        user: User,
        session_id: str,
    ) -> dict[str, Any]:
        contract = self._resolve_contract(db, payload=payload, user=user, session_id=session_id)
        version = db.get(ContractVersion, contract.current_authoritative_version_id) if contract.current_authoritative_version_id else None
        if version is None:
            raise HTTPException(status.HTTP_404_NOT_FOUND, "Current contract version not found")
        contract_file = db.get(ContractFile, version.contract_file_id)
        if contract_file is None:
            raise HTTPException(status.HTTP_404_NOT_FOUND, "Contract file not found")
        source_snapshot = db.get(ContractTextSnapshot, version.text_snapshot_id) if version.text_snapshot_id else None
        proposed_text = _assistant_edit_text(
            source_text=source_snapshot.text if source_snapshot else "",
            instructions=payload.instructions,
        )
        edit_docx = _build_edit_docx(
            contract_title=contract.title,
            base_version_number=version.version_number,
            instructions=payload.instructions,
            source_text=source_snapshot.text if source_snapshot else "",
        )
        storage_object = _store_docx(
            db,
            org_id=user.org_id,
            user_id=user.id,
            filename=f"{_safe_filename(contract.title)}-assistant-edit-v{version.version_number}.docx",
            content=edit_docx,
        )
        edit = ContractEdit(
            org_id=user.org_id,
            contract_id=contract.id,
            contract_version_id=version.id,
            edit_type="assistant_instruction",
            status="proposed",
            original_text=(source_snapshot.text[:4000] if source_snapshot else None),
            replacement_text=proposed_text,
            rationale=payload.instructions,
            citation=[],
            created_by_user_id=user.id,
            updated_by_user_id=user.id,
        )
        db.add(edit)
        db.flush()
        edit_version = ContractVersion(
            org_id=user.org_id,
            contract_id=contract.id,
            contract_file_id=contract_file.id,
            version_number=next_version_number(db, contract_file.id),
            storage_object_id=storage_object.id,
            source=ContractVersionSource.ASSISTANT_EDIT,
            change_summary=f"Assistant edit proposal: {payload.instructions[:240]}",
            is_authoritative=False,
            created_by_user_id=user.id,
            updated_by_user_id=user.id,
        )
        db.add(edit_version)
        db.flush()
        edit.citation = [
            {
                "type": "assistant_edit_version",
                "contract_version_id": edit_version.id,
            }
        ]
        snapshot = ContractTextSnapshot(
            org_id=user.org_id,
            contract_id=contract.id,
            contract_version_id=edit_version.id,
            extraction_method="assistant_edit_text",
            extraction_quality_score=1.0,
            text=proposed_text,
            page_map=source_snapshot.page_map if source_snapshot else None,
            ocr_provider=source_snapshot.ocr_provider if source_snapshot else None,
            validation_status="complete",
            created_by_user_id=user.id,
            updated_by_user_id=user.id,
        )
        db.add(snapshot)
        db.flush()
        edit_version.text_snapshot_id = snapshot.id
        # The proposed edit is NOT authoritative and must not become the file's
        # current version until it is explicitly accepted via accept_contract_edit.
        write_audit_log(
            db,
            action="assistant.contract_edit_created",
            resource_type="contract_edit",
            resource_id=edit.id,
            org_id=user.org_id,
            actor_user_id=user.id,
            after={
                "contract_id": contract.id,
                "base_version_id": version.id,
                "contract_version_id": edit_version.id,
            },
        )
        write_timeline_event(
            db,
            org_id=user.org_id,
            resource_type="contract",
            resource_id=contract.id,
            event_type="assistant.tracked_change_created",
            title="Assistant tracked-change version created",
            actor_user_id=user.id,
            details={"contract_edit_id": edit.id, "contract_version_id": edit_version.id},
        )
        db.flush()
        return {
            "status": "created",
            "artifact_type": "assistant_edit",
            "contract_id": contract.id,
            "base_version_id": version.id,
            "contract_version_id": edit_version.id,
            "contract_edit_id": edit.id,
        }

    def _replicate_contract_version(
        self,
        db: Session,
        *,
        payload: ContractHandleInput,
        user: User,
        session_id: str,
    ) -> dict[str, Any]:
        contract = self._resolve_contract(db, payload=payload, user=user, session_id=session_id)
        version = db.get(ContractVersion, contract.current_authoritative_version_id) if contract.current_authoritative_version_id else None
        if version is None:
            raise HTTPException(status.HTTP_404_NOT_FOUND, "Current contract version not found")
        contract_file = db.get(ContractFile, version.contract_file_id)
        if contract_file is None:
            raise HTTPException(status.HTTP_404_NOT_FOUND, "Contract file not found")
        replicated = ContractVersion(
            org_id=user.org_id,
            contract_id=contract.id,
            contract_file_id=contract_file.id,
            version_number=next_version_number(db, contract_file.id),
            storage_object_id=version.storage_object_id,
            source=ContractVersionSource.ASSISTANT_GENERATED,
            change_summary=f"Replicated from version {version.version_number}",
            is_authoritative=False,
            created_by_user_id=user.id,
            updated_by_user_id=user.id,
        )
        db.add(replicated)
        db.flush()
        source_snapshot = db.get(ContractTextSnapshot, version.text_snapshot_id) if version.text_snapshot_id else None
        if source_snapshot is not None:
            snapshot = ContractTextSnapshot(
                org_id=user.org_id,
                contract_id=contract.id,
                contract_version_id=replicated.id,
                extraction_method=source_snapshot.extraction_method,
                extraction_quality_score=source_snapshot.extraction_quality_score,
                text=source_snapshot.text,
                page_map=source_snapshot.page_map,
                ocr_provider=source_snapshot.ocr_provider,
                validation_status=source_snapshot.validation_status,
                created_by_user_id=user.id,
                updated_by_user_id=user.id,
            )
            db.add(snapshot)
            db.flush()
            replicated.text_snapshot_id = snapshot.id
        write_audit_log(
            db,
            action="assistant.contract_version_replicated",
            resource_type="contract",
            resource_id=contract.id,
            org_id=user.org_id,
            actor_user_id=user.id,
            after={
                "source_version_id": version.id,
                "contract_version_id": replicated.id,
                "contract_file_id": contract_file.id,
            },
        )
        write_timeline_event(
            db,
            org_id=user.org_id,
            resource_type="contract",
            resource_id=contract.id,
            event_type="assistant.contract_version_replicated",
            title="Assistant replicated contract version",
            actor_user_id=user.id,
            details={"source_version_id": version.id, "contract_version_id": replicated.id},
        )
        db.flush()
        return {
            "status": "created",
            "contract_id": contract.id,
            "source_version_id": version.id,
            "contract_version_id": replicated.id,
        }

    def _run_playbook_review(
        self,
        db: Session,
        *,
        payload: PlaybookToolInput,
        user: User,
        session_id: str,
        create_redline: bool,
    ) -> dict[str, Any]:
        if not has_permission(user.permission_values, "contract:read"):
            raise HTTPException(status.HTTP_403_FORBIDDEN, "Missing permission: contract:read")
        if not has_permission(user.permission_values, "playbook:run"):
            raise HTTPException(status.HTTP_403_FORBIDDEN, "Missing permission: playbook:run")
        if create_redline and not has_permission(user.permission_values, "contract:redline"):
            raise HTTPException(status.HTTP_403_FORBIDDEN, "Missing permission: contract:redline")
        contract = self._resolve_contract(db, payload=payload, user=user, session_id=session_id)
        playbook = get_playbook_for_user(db, playbook_id=payload.playbook_id, user=user)
        version = select_run_version(
            db,
            playbook=playbook,
            org_id=user.org_id,
            version_id=payload.playbook_version_id,
            test_mode=payload.test_mode,
        )
        artifacts = execute_playbook_run(
            db,
            user=user,
            playbook=playbook,
            version=version,
            contract=contract,
            create_redline=create_redline,
        )
        db.flush()
        result = {
            "status": artifacts.run.status,
            "artifact_type": "playbook_redline" if artifacts.redline_version else "playbook_review",
            "contract_id": contract.id,
            "playbook_id": playbook.id,
            "playbook_version_id": version.id,
            "playbook_name": playbook.name,
            "playbook_run_id": artifacts.run.id,
            "deviation_count": len(artifacts.deviations),
            "deviations": [
                {
                    "playbook_deviation_id": row.id,
                    "clause_type": row.clause_type,
                    "severity": row.severity,
                    "issue": row.issue,
                    "suggested_fix": row.suggested_fix,
                }
                for row in artifacts.deviations
            ],
        }
        if artifacts.redline_version is not None:
            result.update(
                {
                    "base_version_id": artifacts.contract_edit.contract_version_id if artifacts.contract_edit else None,
                    "contract_version_id": artifacts.redline_version.id,
                    "contract_edit_id": artifacts.contract_edit.id if artifacts.contract_edit else None,
                }
            )
        return result

    def _ask_contract_brain(
        self,
        db: Session,
        *,
        payload: BrainAskInput,
        user: User,
        session_id: str,
    ) -> dict[str, Any]:
        contract_id = payload.contract_id
        if payload.contract_handle or (payload.query_scope == "contract" and contract_id):
            contract = self._resolve_contract(db, payload=payload, user=user, session_id=session_id)
            contract_id = contract.id
        if payload.query_scope == "contract" and not contract_id:
            raise HTTPException(status.HTTP_422_UNPROCESSABLE_ENTITY, "contract handle or contract_id is required")
        if payload.project_id:
            get_project_for_user(db, project_id=payload.project_id, user=user)
        parsed = BrainQueryParseOutput(query_scope=payload.query_scope)
        context = assemble_context(
            db,
            user=user,
            question=payload.question,
            scope=payload.query_scope,
            contract_id=contract_id,
            project_id=payload.project_id,
            parsed=parsed,
        )
        return {
            "status": "retrieved",
            "scope": payload.query_scope,
            "source_count": context["source_count"],
            "graph_fact_count": len(context["graph_facts"]),
            "vector_chunk_count": len(context["vector_chunks"]),
            "fulltext_clause_count": len(context["fulltext_clauses"]),
            "context_text": context["context_text"][:12000],
            "context_truncated": len(context["context_text"]) > 12000,
        }

    async def _submit_for_approval(
        self,
        db: Session,
        *,
        payload: ApprovalSubmitInput,
        user: User,
        session_id: str,
    ) -> dict[str, Any]:
        contract = self._resolve_contract(db, payload=payload, user=user, session_id=session_id)
        requests = await submit_contract_for_approval(
            db,
            user=user,
            contract=contract,
            contract_version_id=contract.current_authoritative_version_id,
            approver_user_id=payload.approver_user_id,
            approver_role=payload.approver_role,
        )
        db.flush()
        return {
            "status": "submitted",
            "contract_id": contract.id,
            "approval_request_ids": [approval.id for approval in requests],
            "approval_count": len(requests),
        }

    async def _send_for_signature(
        self,
        db: Session,
        *,
        payload: SignatureSendInput,
        user: User,
        session_id: str,
    ) -> dict[str, Any]:
        contract = self._resolve_contract(db, payload=payload, user=user, session_id=session_id)
        if payload.override_lifecycle and not has_permission(
            user.permission_values, "contract:lifecycle_override"
        ):
            raise HTTPException(
                status.HTTP_403_FORBIDDEN,
                "Overriding the approved-before-signature gate requires contract:lifecycle_override",
            )
        if contract.lifecycle_stage != ContractLifecycleStage.APPROVED and not payload.override_lifecycle:
            raise HTTPException(status.HTTP_409_CONFLICT, "Contract must be approved before signature")
        version_id = contract.current_authoritative_version_id
        version = db.get(ContractVersion, version_id)
        if version is None or version.org_id != user.org_id:
            raise HTTPException(status.HTTP_404_NOT_FOUND, "Contract version not found")
        storage_object = db.get(StorageObject, version.storage_object_id)
        if storage_object is None or storage_object.org_id != user.org_id:
            raise HTTPException(status.HTTP_404_NOT_FOUND, "Stored file not found")
        validate_signature_recipients(
            db, contract=contract, org_id=user.org_id, recipients=payload.recipients
        )
        content = storage_service.read_bytes(storage_object.storage_key)
        envelope = await docusign_client.create_envelope(
            filename=storage_object.filename,
            recipients=[recipient.model_dump(mode="json") for recipient in payload.recipients],
            content=content,
        )
        signature = SignatureRequest(
            org_id=user.org_id,
            contract_id=contract.id,
            contract_version_id=version.id,
            provider_envelope_id=envelope.envelope_id,
            status=SignatureStatus.SENT,
            sent_by_user_id=user.id,
            sent_at=datetime.now(UTC),
            metadata_json=envelope.metadata,
            created_by_user_id=user.id,
            updated_by_user_id=user.id,
        )
        db.add(signature)
        db.flush()
        for index, recipient in enumerate(payload.recipients, start=1):
            db.add(
                SignatureRecipient(
                    org_id=user.org_id,
                    signature_request_id=signature.id,
                    name=recipient.name,
                    email=str(recipient.email),
                    role=recipient.role,
                    routing_order=str(index),
                    status="sent",
                    created_by_user_id=user.id,
                    updated_by_user_id=user.id,
                )
            )
            try:
                await resend_client.send_email(
                    to=str(recipient.email),
                    subject=f"Signature requested: {contract.title}",
                    html=(
                        f"<p>You have been requested to sign <b>{contract.title}</b>.</p>"
                        f"<p>Envelope: {envelope.envelope_id or 'mock'}</p>"
                    ),
                )
            except Exception as exc:
                logging.getLogger(__name__).warning(
                    "assistant signature notification email failed for %s: %s",
                    recipient.email,
                    exc,
                )
        transition_contract_stage(
            db,
            contract=contract,
            to_stage=ContractLifecycleStage.SIGNATURE_PENDING,
            actor_user_id=user.id,
            reason="Sent for signature by assistant",
            override=payload.override_lifecycle,
            override_authorized=payload.override_lifecycle,
        )
        db.flush()
        return {
            "status": "sent",
            "contract_id": contract.id,
            "signature_request_id": signature.id,
            "recipient_count": len(payload.recipients),
            "provider_status": envelope.status,
        }

    def _extract_obligations(
        self,
        db: Session,
        *,
        payload: ExtractObligationsInput,
        user: User,
        session_id: str,
    ) -> dict[str, Any]:
        contract = self._resolve_contract(db, payload=payload, user=user, session_id=session_id)
        version = db.get(ContractVersion, contract.current_authoritative_version_id) if contract.current_authoritative_version_id else None
        snapshot = db.get(ContractTextSnapshot, version.text_snapshot_id) if version and version.text_snapshot_id else None
        if version is None or snapshot is None:
            raise HTTPException(status.HTTP_409_CONFLICT, "Contract has no extractable authoritative version")
        job = create_job(
            db,
            org_id=user.org_id,
            job_type="obligation_extraction",
            resource_type="contract",
            resource_id=contract.id,
            created_by_user_id=user.id,
            idempotency_key=f"obligation_extraction:{version.id}:{snapshot.id}:assistant:{utcnow().timestamp()}",
            metadata={"contract_version_id": version.id, "text_snapshot_id": snapshot.id},
        )
        db.flush()
        dispatch_job(db, job=job)
        db.flush()
        return {"status": "queued", "contract_id": contract.id, "job_id": job.id, "job_type": job.job_type}

    def _create_tabular_review(
        self,
        db: Session,
        *,
        payload: TabularReviewInput,
        user: User,
        session_id: str,
    ) -> dict[str, Any]:
        contract_ids = list(dict.fromkeys(payload.contract_ids))
        for handle in payload.contract_handles:
            contract = self._resolve_contract(
                db,
                payload=ContractHandleInput(contract_handle=handle),
                user=user,
                session_id=session_id,
            )
            contract_ids.append(contract.id)
        contract_ids = list(dict.fromkeys(contract_ids))
        if payload.project_id:
            get_project_for_user(db, project_id=payload.project_id, user=user)
            if not contract_ids:
                contract_ids = list(
                    db.scalars(
                        select(ProjectContract.contract_id).where(
                            ProjectContract.org_id == user.org_id,
                            ProjectContract.project_id == payload.project_id,
                        )
                    ).all()
                )
        if not contract_ids:
            raise HTTPException(status.HTTP_422_UNPROCESSABLE_ENTITY, "No contracts selected")
        for contract_id in contract_ids:
            get_contract_for_user(db, contract_id=contract_id, user=user)
        review = TabularReview(
            org_id=user.org_id,
            name=payload.name,
            project_id=payload.project_id,
            source_contract_ids=contract_ids,
            status="running",
            created_by_user_id=user.id,
            updated_by_user_id=user.id,
        )
        db.add(review)
        db.flush()
        columns = []
        for position, column_payload in enumerate(payload.columns):
            column = TabularReviewColumn(
                org_id=user.org_id,
                tabular_review_id=review.id,
                name=column_payload.name,
                prompt=column_payload.prompt,
                position=position,
                created_by_user_id=user.id,
                updated_by_user_id=user.id,
            )
            db.add(column)
            db.flush()
            columns.append(column)
        cells = []
        for contract_id in contract_ids:
            for column in columns:
                cell = TabularReviewCell(
                    org_id=user.org_id,
                    tabular_review_id=review.id,
                    column_id=column.id,
                    contract_id=contract_id,
                    status=TabularCellStatus.PENDING,
                    created_by_user_id=user.id,
                    updated_by_user_id=user.id,
                )
                db.add(cell)
                db.flush()
                cells.append(cell)
        dispatch_cells(db, user=user, review=review, cells=cells)
        return {
            "status": "created",
            "tabular_review_id": review.id,
            "contract_count": len(contract_ids),
            "column_count": len(columns),
            "cell_count": len(cells),
        }

    def _read_table_cells(
        self,
        db: Session,
        *,
        payload: ReadTableCellsInput,
        user: User,
    ) -> dict[str, Any]:
        review = db.get(TabularReview, payload.tabular_review_id)
        if review is None or review.org_id != user.org_id or review.deleted_at is not None:
            raise HTTPException(status.HTTP_404_NOT_FOUND, "Tabular review not found")
        for contract_id in review.source_contract_ids or []:
            get_contract_for_user(db, contract_id=contract_id, user=user)
        columns = db.scalars(
            select(TabularReviewColumn)
            .where(TabularReviewColumn.tabular_review_id == review.id)
            .order_by(TabularReviewColumn.position.asc())
        ).all()
        cells = db.scalars(
            select(TabularReviewCell).where(
                TabularReviewCell.org_id == user.org_id,
                TabularReviewCell.tabular_review_id == review.id,
            )
        ).all()
        return {
            "tabular_review_id": review.id,
            "columns": [{"column_id": column.id, "name": column.name} for column in columns],
            "cells": [
                {
                    "column_id": cell.column_id,
                    "contract_id": cell.contract_id,
                    "status": cell.status,
                    "answer": cell.answer,
                    "confidence": cell.confidence,
                    "citation_count": len(cell.citations or []),
                }
                for cell in cells
            ],
        }

    def _external_share(
        self,
        db: Session,
        *,
        payload: ExternalShareInput,
        user: User,
        session_id: str,
    ) -> dict[str, Any]:
        contract = self._resolve_contract(db, payload=payload, user=user, session_id=session_id)
        version = db.get(ContractVersion, contract.current_authoritative_version_id) if contract.current_authoritative_version_id else None
        token = secrets.token_urlsafe(32)
        expires_at = (
            datetime.now(UTC) + timedelta(days=payload.expires_in_days)
            if payload.expires_in_days
            else None
        )
        share = ContractShare(
            org_id=user.org_id,
            contract_id=contract.id,
            contract_version_id=version.id if version else None,
            token_hash=_hash_secret(token),
            passcode_hash=_hash_secret(payload.passcode) if payload.passcode else None,
            access_mode=ShareAccessMode.DOWNLOAD_ALLOWED if payload.download_allowed else ShareAccessMode.VIEW_ONLY,
            expires_at=expires_at,
            download_allowed=payload.download_allowed,
            created_by_user_id=user.id,
            updated_by_user_id=user.id,
        )
        db.add(share)
        db.flush()
        write_audit_log(
            db,
            action="contract.share_created",
            resource_type="contract_share",
            resource_id=share.id,
            org_id=user.org_id,
            actor_user_id=user.id,
            after={
                "contract_id": contract.id,
                "contract_version_id": share.contract_version_id,
                "download_allowed": share.download_allowed,
                "expires_at": share.expires_at.isoformat() if share.expires_at else None,
                "created_by": "assistant_tool",
            },
        )
        write_timeline_event(
            db,
            org_id=user.org_id,
            resource_type="contract",
            resource_id=contract.id,
            event_type="contract.share_created",
            title="Contract share created by assistant",
            actor_user_id=user.id,
            details={"share_id": share.id, "download_allowed": share.download_allowed},
        )
        return {"status": "created", "contract_id": contract.id, "contract_share_id": share.id, "external_share_token": token}

    def _archive_contract(
        self,
        db: Session,
        *,
        payload: ArchiveContractInput,
        user: User,
        session_id: str,
    ) -> dict[str, Any]:
        contract = self._resolve_contract(db, payload=payload, user=user, session_id=session_id)
        transition_contract_stage(
            db,
            contract=contract,
            to_stage=ContractLifecycleStage.ARCHIVED,
            actor_user_id=user.id,
            reason=payload.reason or "Archived by assistant",
            override=True,
            override_authorized=True,
        )
        db.flush()
        return {"status": "archived", "contract_id": contract.id, "lifecycle_stage": contract.lifecycle_stage}

    def _resolve_contract(
        self,
        db: Session,
        *,
        payload: ContractHandleInput,
        user: User,
        session_id: str,
    ) -> Contract:
        contract_id = payload.contract_id
        if payload.contract_handle:
            handle = db.scalar(
                select(AssistantContractHandle).where(
                    AssistantContractHandle.org_id == user.org_id,
                    AssistantContractHandle.session_id == session_id,
                    AssistantContractHandle.handle == payload.contract_handle,
                )
            )
            if handle is None:
                raise HTTPException(status.HTTP_404_NOT_FOUND, "Contract handle not found")
            contract_id = handle.contract_id
        if not contract_id:
            raise HTTPException(status.HTTP_422_UNPROCESSABLE_ENTITY, "contract_id or contract_handle is required")
        return get_contract_for_user(db, contract_id=contract_id, user=user)


def _idempotency_key(tool_name: str, session_id: str, payload: dict[str, Any]) -> str:
    canonical = json.dumps(payload, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(f"{tool_name}:{session_id}:{canonical}".encode("utf-8")).hexdigest()


def _hash_secret(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def _store_docx(db: Session, *, org_id: str, user_id: str, filename: str, content: bytes) -> StorageObject:
    stored = storage_service.save_bytes(
        org_id=org_id,
        filename=filename,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        content=content,
    )
    storage_object = StorageObject(
        org_id=org_id,
        storage_key=stored.storage_key,
        filename=stored.filename,
        mime_type=stored.mime_type,
        size_bytes=stored.size_bytes,
        sha256_hash=stored.sha256_hash,
        storage_backend=StorageBackend.LOCAL_VOLUME,
        created_by_user_id=user_id,
        updated_by_user_id=user_id,
    )
    db.add(storage_object)
    db.flush()
    return storage_object


def _build_generated_contract_docx(*, title: str, instructions: str) -> tuple[str, bytes]:
    from docx import Document

    sections = [
        ("Parties", "Identify the contracting parties and fill in any missing legal names."),
        ("Purpose", "Describe the commercial purpose of this contract."),
        ("Key Terms", instructions),
        ("Obligations", "List each party's main obligations clearly and separately."),
        ("Payment and Fees", "State payment amounts, timing, taxes, and invoicing rules if applicable."),
        ("Confidentiality", "Include confidentiality duties, exclusions, and survival period."),
        ("Term and Termination", "State effective date, renewal, termination rights, and notice periods."),
        ("Governing Law", "State governing law and dispute forum."),
        ("Open Issues", "Confirm all bracketed or business-specific terms before signature."),
    ]
    text_parts = [title]
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph("Assistant-generated working draft. Review before use.")
    for heading, body in sections:
        document.add_heading(heading, level=2)
        document.add_paragraph(body)
        text_parts.extend([heading, body])
    buffer = BytesIO()
    document.save(buffer)
    return "\n\n".join(text_parts), buffer.getvalue()


def _build_edit_docx(
    *,
    contract_title: str,
    base_version_number: int,
    instructions: str,
    source_text: str,
) -> bytes:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = Document()
    _enable_word_track_revisions(document, OxmlElement=OxmlElement, qn=qn)
    document.add_heading(f"{contract_title} - Assistant Redline Proposal", level=1)
    document.add_paragraph(f"Base version: V{base_version_number}")
    document.add_heading("Requested Change", level=2)
    document.add_paragraph(instructions)
    document.add_heading("Native Word Tracked Changes", level=2)
    document.add_paragraph(
        "This version contains native Word revision markup. "
        "Use accept/reject endpoints to make the proposal authoritative or close it."
    )
    revision_paragraph = document.add_paragraph()
    _append_deleted_text(
        revision_paragraph,
        source_text or "No source text snapshot was available.",
        author="Legal AI Assistant",
        revision_id="1",
        OxmlElement=OxmlElement,
        qn=qn,
    )
    _append_inserted_text(
        revision_paragraph,
        _assistant_edit_text(source_text=source_text, instructions=instructions),
        author="Legal AI Assistant",
        revision_id="2",
        OxmlElement=OxmlElement,
        qn=qn,
    )
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _enable_word_track_revisions(document: Any, *, OxmlElement: Any, qn: Any) -> None:
    settings = document.settings.element
    if settings.find(qn("w:trackRevisions")) is None:
        settings.append(OxmlElement("w:trackRevisions"))


def _append_deleted_text(
    paragraph: Any,
    text: str,
    *,
    author: str,
    revision_id: str,
    OxmlElement: Any,
    qn: Any,
) -> None:
    deletion = OxmlElement("w:del")
    _set_revision_attrs(deletion, author=author, revision_id=revision_id, qn=qn)
    _append_revision_runs(deletion, text, text_tag="w:delText", OxmlElement=OxmlElement)
    paragraph._p.append(deletion)


def _append_inserted_text(
    paragraph: Any,
    text: str,
    *,
    author: str,
    revision_id: str,
    OxmlElement: Any,
    qn: Any,
) -> None:
    insertion = OxmlElement("w:ins")
    _set_revision_attrs(insertion, author=author, revision_id=revision_id, qn=qn)
    _append_revision_runs(insertion, text, text_tag="w:t", OxmlElement=OxmlElement)
    paragraph._p.append(insertion)


def _set_revision_attrs(element: Any, *, author: str, revision_id: str, qn: Any) -> None:
    element.set(qn("w:id"), revision_id)
    element.set(qn("w:author"), author)
    element.set(qn("w:date"), utcnow().replace(microsecond=0).isoformat())


def _append_revision_runs(parent: Any, text: str, *, text_tag: str, OxmlElement: Any) -> None:
    xml_space = "{http://www.w3.org/XML/1998/namespace}space"
    for line_index, line in enumerate(text.splitlines() or [""]):
        if line_index:
            break_run = OxmlElement("w:r")
            break_run.append(OxmlElement("w:br"))
            parent.append(break_run)
        run = OxmlElement("w:r")
        text_element = OxmlElement(text_tag)
        text_element.set(xml_space, "preserve")
        text_element.text = line
        run.append(text_element)
        parent.append(run)


def _assistant_edit_text(*, source_text: str, instructions: str) -> str:
    return "\n\n".join(
        [
            source_text,
            "[Assistant proposed tracked change]",
            instructions,
        ]
    ).strip()


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", value).strip("-._")
    return cleaned[:120] or "contract"


tool_runtime = ToolRuntime()
